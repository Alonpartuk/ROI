from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('BigQuery Views Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('HubSpot CEO Metrics Suite - Complete Reference')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Overview section
doc.add_heading('Overview', level=1)
doc.add_paragraph('This document describes all 44 BigQuery views created for tracking sales performance, SDR activity, deal risk analysis, and AI-powered executive summaries. All views are in the octup-testing.hubspot_data dataset.')

doc.add_heading('Data Sources', level=2)
doc.add_paragraph('- deals_snapshots: Daily snapshot of all deals in "3PL New Business" pipeline')
doc.add_paragraph('- meetings_snapshots: Daily snapshot of all HubSpot meetings')

doc.add_heading('Key Business Rules', level=2)
doc.add_paragraph('- Pipeline Filter: Most views filter for "3PL New Business" pipeline')
doc.add_paragraph('- SDR Tracking: Uses created_by_name with fallback to owner_name')
doc.add_paragraph('- Week Definition: Weeks start on Monday (using WEEK(MONDAY))')
doc.add_paragraph('- ARR Value: Uses hs_arr field with fallback to amount')

# Table of Contents
doc.add_page_break()
doc.add_heading('Table of Contents', level=1)

toc_items = [
    ('Base/Core Views (31)', [
        'v_latest_snapshot', 'v_pipeline_summary', 'v_deal_aging', 'v_deal_aging_summary',
        'v_avg_time_in_stage', 'v_pipeline_concentration', 'v_sales_velocity', 'v_stage_conversion',
        'v_win_rate_analysis', 'v_forecast_weighted', 'v_close_date_slippage', 'v_slippage_summary',
        'v_owner_leaderboard', 'v_multi_threading', 'v_next_step_coverage', 'v_lost_deal_analysis',
        'v_won_deal_analysis', 'v_pipeline_trend', 'v_deal_type_analysis', 'v_closing_this_month',
        'v_stalled_deals', 'v_new_deals_daily', 'v_pipeline_coverage', 'v_forecast_distribution',
        'v_weekly_comparison', 'v_priority_analysis', 'v_stage_leakage', 'v_ceo_dashboard',
        'v_contact_analysis', 'v_company_analysis', 'v_full_deal_details'
    ]),
    ('SDR Activity Views (3)', ['v_sdr_activity_weekly', 'v_sdr_meeting_outcomes', 'v_sdr_leaderboard']),
    ('Deal Risk Views (5)', ['v_deals_at_risk', 'v_orphaned_deals_summary', 'v_daily_deal_movements', 'v_stalled_deals_alert', 'v_at_risk_by_owner']),
    ('SDR Forecast View (1)', ['v_sdr_forecast_model']),
    ('AI Summary Views (4)', ['v_ai_executive_summary', 'v_ai_insight_simple', 'v_weekly_summary_data', 'v_ceo_ai_summary'])
]

for category, views in toc_items:
    doc.add_heading(category, level=2)
    for view in views:
        doc.add_paragraph(f'  - {view}')

# Base Views
doc.add_page_break()
doc.add_heading('Base/Core Views', level=1)

views_data = [
    ('v_latest_snapshot', 'Get the most recent snapshot of all deals', 'Filters to only the latest snapshot_date in the deals table.', 'All columns from deals_snapshots for the current day'),
    ('v_pipeline_summary', 'High-level dashboard of pipeline metrics', 'Aggregates deal counts, values, and averages grouped by snapshot date and pipeline.', 'snapshot_date, pipeline_label, total_deals, open_deals, won_deals, lost_deals, open_pipeline_value, weighted_pipeline_value'),
    ('v_deal_aging', 'SLA-based RAG status for deal aging', 'Shows open deals with aging status and priority score based on deal_age_status field.', 'hs_object_id, dealname, deal_age_status, age_priority_score, days_in_current_stage'),
    ('v_deal_aging_summary', 'Summary statistics by deal age status', 'Groups open deals by deal_age_status with counts and values.', 'deal_age_status, deal_count, total_value, avg_days_in_stage, pct_of_deals'),
    ('v_avg_time_in_stage', 'Time deals spend in each stage', 'Calculates avg, median, min, max per pipeline stage.', 'dealstage_label, deals_in_stage, avg_days, median_days, min_days, max_days'),
    ('v_pipeline_concentration', 'Concentration risk analysis', 'Calculates % of pipeline value in each stage. Flags >50% as High Risk.', 'dealstage_label, stage_value, pct_of_pipeline_value, concentration_risk'),
    ('v_sales_velocity', 'Sales velocity: V = (n x L x W) / T', 'n=opportunities, L=avg deal value, W=win rate, T=sales cycle length.', 'owner_name, num_opportunities, avg_deal_value, win_rate_pct, sales_velocity_monthly'),
    ('v_stage_conversion', 'Stage-to-stage transitions with deal details', 'Shows each deal that moved stages with direct HubSpot link.', 'deal_id, deal_name, hubspot_link, from_stage, to_stage, deal_value'),
    ('v_win_rate_analysis', 'Win rate analysis by owner', 'Calculates won/(won+lost) ratio per owner.', 'owner_name, won_count, lost_count, win_rate_pct, avg_won_deal_size'),
    ('v_forecast_weighted', 'Weighted pipeline by forecast category', 'Groups by HubSpot forecast category with weighted amounts.', 'hs_forecast_category, deal_count, total_amount, weighted_forecast'),
    ('v_close_date_slippage', 'Track when close dates are pushed out', 'Compares consecutive days to detect closedate changes.', 'dealname, prev_closedate, curr_closedate, days_slipped, slippage_category'),
    ('v_slippage_summary', 'Summary of slippage rates', 'Calculates % of deals that slipped their close date.', 'total_open_deals, slipped_count, slippage_rate_pct, value_slippage_rate_pct'),
    ('v_owner_leaderboard', 'Sales rep performance rankings', 'Aggregates all key metrics per owner for comparison.', 'owner_name, open_deals, won_deals, pipeline_value, won_value, win_rate_pct'),
    ('v_multi_threading', 'Contact coverage per deal', 'Groups deals by contact count (4+, 2-3, 1, None).', 'threading_level, deal_count, total_value, win_rate_pct'),
    ('v_next_step_coverage', 'Deals with defined next steps', 'Calculates % of open deals with hs_next_step populated.', 'owner_name, total_open_deals, deals_with_next_step, next_step_coverage_pct'),
    ('v_lost_deal_analysis', 'Analyze reasons for lost deals', 'Groups lost deals by closed_lost_reason.', 'closed_lost_reason, lost_count, lost_value, pct_of_losses'),
    ('v_won_deal_analysis', 'Analyze won deals patterns', 'Groups won deals by closed_won_reason.', 'closed_won_reason, won_count, won_value, pct_of_wins'),
    ('v_pipeline_trend', 'Pipeline changes over time', 'Daily aggregates with LAG for day-over-day changes.', 'snapshot_date, open_pipeline_value, prev_pipeline_value, pipeline_change'),
    ('v_deal_type_analysis', 'Performance by deal type', 'Groups by dealtype field.', 'dealtype, total_deals, pipeline_value, win_rate_pct'),
    ('v_closing_this_month', 'Deals expected to close this month', 'Filters open deals where closedate is in current month.', 'dealname, amount, closedate, days_to_close, deal_age_status'),
    ('v_stalled_deals', 'Deals with no recent activity', 'Calculates days since last activity.', 'dealname, days_since_last_activity, activity_status'),
    ('v_new_deals_daily', 'New deal creation by day', 'Filters deals where createdate matches snapshot_date.', 'created_date, owner_name, new_deals, new_deals_value'),
    ('v_pipeline_coverage', 'Pipeline coverage ratio vs targets', 'Calculates ratio of pipeline to MTD wins (target: 3x).', 'pipeline_value, coverage_ratio, coverage_status'),
    ('v_forecast_distribution', 'Pipeline by forecast category', 'Shows % breakdown by forecast category.', 'hs_forecast_category, deal_count, pct_of_pipeline'),
    ('v_weekly_comparison', 'Week-over-week pipeline comparison', 'Aggregates by week with LAG for comparison.', 'week_start, pipeline_value, prev_week_pipeline, pipeline_change_pct'),
    ('v_priority_analysis', 'Performance by deal priority', 'Groups by hs_priority field.', 'hs_priority, deal_count, pipeline_value'),
    ('v_stage_leakage', 'How deals exit each stage', 'Identifies stage exits (Won/Lost/Moved).', 'from_stage, exit_type, exit_count, exit_value'),
    ('v_ceo_dashboard', 'Executive summary dashboard', 'Single-row view with comprehensive pipeline metrics.', 'total_open_deals, total_pipeline_value, weighted_pipeline_value, win_rate_pct, red_deals_count'),
    ('v_contact_analysis', 'Contact information per deal', 'Shows primary contact details for latest snapshot.', 'dealname, contact_count, primary_contact_name, primary_contact_email'),
    ('v_company_analysis', 'Pipeline analysis by company', 'Groups deals by company with demographics.', 'company_name, company_industry, deal_count, pipeline_value'),
    ('v_full_deal_details', 'Complete deal record', 'Returns all columns from deals_snapshots for current date.', 'All deal fields including contacts and company data')
]

for i, (name, purpose, logic, columns) in enumerate(views_data, 1):
    doc.add_heading(f'{i}. {name}', level=2)
    doc.add_paragraph(f'Purpose: {purpose}')
    doc.add_paragraph(f'Business Logic: {logic}')
    doc.add_paragraph(f'Key Columns: {columns}')
    doc.add_paragraph()

# SDR Activity Views
doc.add_page_break()
doc.add_heading('SDR Activity Views', level=1)

sdr_views = [
    ('v_sdr_activity_weekly', 'Track weekly SDR meeting performance',
     'Counts meetings created each week by SDR. SDR determined by created_by_name with fallback to owner_name.',
     'week_start, sdr_name, meetings_booked_count, meetings_held_count, held_rate_pct, no_show_rate_pct'),
    ('v_sdr_meeting_outcomes', 'Meeting outcomes breakdown by SDR',
     'Pivots meeting outcomes into rows. Shows % of total for each outcome type.',
     'week_start, sdr_name, meeting_outcome, meeting_count, pct_of_total'),
    ('v_sdr_leaderboard', 'Weekly SDR performance ranking',
     'Shows current week SDR performance with at-risk portfolio. Ranks by multiple criteria.',
     'sdr_name, meetings_held_count, held_rate_pct, at_risk_value, rank_by_meetings_held')
]

for i, (name, purpose, logic, columns) in enumerate(sdr_views, 32):
    doc.add_heading(f'{i}. {name}', level=2)
    doc.add_paragraph(f'Purpose: {purpose}')
    doc.add_paragraph(f'Business Logic: {logic}')
    doc.add_paragraph(f'Key Columns: {columns}')
    doc.add_paragraph()

# Deal Risk Views
doc.add_page_break()
doc.add_heading('Deal Risk Views', level=1)

risk_views = [
    ('v_deals_at_risk', 'Flag deals with risk indicators',
     'Applies 5 risk flags: is_unassigned_risk (Kurt/Hanan), is_stalled (>14 days), is_ghosted (no contact >5 days), is_not_3pl_match, is_stalled_delayed (Delayed >30 days).',
     'dealname, arr_value, owner_name, is_at_risk, primary_risk_reason, risk_flag_count'),
    ('v_orphaned_deals_summary', 'Deals owned by placeholder owners',
     'Filters deals where is_unassigned_risk = TRUE. Groups by owner to show total exposure.',
     'owner_name, deal_count, total_arr_at_risk, top_deals'),
    ('v_daily_deal_movements', 'Track all deal stage transitions',
     'Compares each day snapshot to previous day. Identifies stage changes using LAG().',
     'deal_name, hubspot_link, previous_stage, current_stage, movement_type, movement_description'),
    ('v_stalled_deals_alert', 'Prioritized stalled deals list',
     'Shows deals stalled >14 days or in Delayed >30 days. Assigns priority rank and recommended action.',
     'deal_name, value_arr, owner_name, days_in_current_stage, recommended_action, priority_rank'),
    ('v_at_risk_by_owner', 'At-risk ARR by deal owner',
     'Aggregates risk metrics per owner. Shows count and ARR value exposure.',
     'owner_name, total_arr, at_risk_deals, at_risk_value, at_risk_arr_pct')
]

for i, (name, purpose, logic, columns) in enumerate(risk_views, 35):
    doc.add_heading(f'{i}. {name}', level=2)
    doc.add_paragraph(f'Purpose: {purpose}')
    doc.add_paragraph(f'Business Logic: {logic}')
    doc.add_paragraph(f'Key Columns: {columns}')
    doc.add_paragraph()

# SDR Forecast
doc.add_page_break()
doc.add_heading('SDR Forecast View', level=1)

doc.add_heading('40. v_sdr_forecast_model', level=2)
doc.add_paragraph('Purpose: Forecast SDR performance and calculate gap to $1.6M ARR target.')
doc.add_paragraph('Business Logic:')
doc.add_paragraph('1. Calculate 4-week rolling average of meetings per SDR')
doc.add_paragraph('2. Get historical win rate and average deal size from closed deals')
doc.add_paragraph('3. Project next 4 weeks of meetings and expected ARR')
doc.add_paragraph('4. Allocate $1.6M target proportionally based on performance')
doc.add_paragraph('5. Calculate gap between forecast and target')
doc.add_paragraph()
doc.add_paragraph('Formula: Expected ARR = Forecasted Held Meetings x Win Rate x Avg Deal Size')
doc.add_paragraph()
doc.add_paragraph('Key Columns: sdr_name, avg_weekly_held, expected_arr_contribution, sdr_arr_target, arr_gap_to_target, forecast_status')

# AI Summary Views
doc.add_page_break()
doc.add_heading('AI Summary Views', level=1)

ai_views = [
    ('v_ai_executive_summary', 'AI-generated executive summary using Gemini 1.5 Pro',
     'Aggregates KPIs, constructs dynamic prompt, calls Gemini via BigQuery ML. Returns 3-bullet CEO summary.',
     'total_pipeline_value, weighted_value, count_at_risk, executive_insight, generated_at'),
    ('v_ai_insight_simple', 'Simplified AI insight view',
     'Wraps v_ai_executive_summary to return just the insight for Looker Studio.',
     'executive_insight, generated_at'),
    ('v_weekly_summary_data', 'Aggregated weekly data for AI',
     'Compares current vs 7-days-ago snapshot. Calculates week-over-week changes.',
     'current_pipeline_value, pipeline_change, pipeline_change_pct, top_owners_summary'),
    ('v_ceo_ai_summary', 'Latest AI-generated CEO summary',
     'Calls fn_generate_ceo_summary() which uses gemini_pro_model.',
     'generated_at, summary')
]

for i, (name, purpose, logic, columns) in enumerate(ai_views, 41):
    doc.add_heading(f'{i}. {name}', level=2)
    doc.add_paragraph(f'Purpose: {purpose}')
    doc.add_paragraph(f'Business Logic: {logic}')
    doc.add_paragraph(f'Key Columns: {columns}')
    doc.add_paragraph()

# Known Limitations
doc.add_page_break()
doc.add_heading('Known Limitations', level=1)

doc.add_heading('SDR Attribution Issue', level=2)
doc.add_paragraph('Problem: Meetings booked by SDRs (like Chanan) are attributed to the AE whose calendar link was used (Blake, Jay).')
doc.add_paragraph()
doc.add_paragraph('Current Behavior:')
doc.add_paragraph('- HubSpot hs_created_by_user_id shows who created the CRM record, not who initiated the booking')
doc.add_paragraph('- When SDR books via AE calendar link, meeting is attributed to AE')
doc.add_paragraph()
doc.add_paragraph('Workaround Options:')
doc.add_paragraph('1. Track by deal ownership: Count meetings on deals owned by the SDR')
doc.add_paragraph('2. Custom HubSpot property: Add "Booked By SDR" field')
doc.add_paragraph('3. Process change: Have SDRs create meetings under their own accounts')

# Refresh Schedule
doc.add_heading('Refresh Schedule', level=1)
doc.add_paragraph('All views read from snapshot tables that are refreshed by the ETL pipeline:')
doc.add_paragraph('- Deals: Daily snapshot at ETL run time')
doc.add_paragraph('- Meetings: Daily snapshot at ETL run time')
doc.add_paragraph()
doc.add_paragraph('Views are always based on the latest snapshot_date available.')

# Quick Reference
doc.add_page_break()
doc.add_heading('Quick Reference: Views by Category', level=1)

categories = [
    ('Pipeline Overview', 'v_pipeline_summary, v_pipeline_trend, v_pipeline_coverage, v_pipeline_concentration'),
    ('Deal Health', 'v_deal_aging, v_deal_aging_summary, v_stalled_deals, v_stalled_deals_alert'),
    ('Risk Analysis', 'v_deals_at_risk, v_at_risk_by_owner, v_orphaned_deals_summary'),
    ('Performance', 'v_owner_leaderboard, v_win_rate_analysis, v_sales_velocity'),
    ('SDR Metrics', 'v_sdr_activity_weekly, v_sdr_meeting_outcomes, v_sdr_leaderboard, v_sdr_forecast_model'),
    ('Forecasting', 'v_forecast_weighted, v_forecast_distribution, v_closing_this_month'),
    ('Movement Tracking', 'v_daily_deal_movements, v_stage_conversion, v_stage_leakage'),
    ('Slippage', 'v_close_date_slippage, v_slippage_summary'),
    ('Contact/Company', 'v_contact_analysis, v_company_analysis, v_multi_threading'),
    ('Executive', 'v_ceo_dashboard, v_ai_executive_summary, v_ceo_ai_summary')
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Category'
hdr_cells[1].text = 'Views'

for category, views in categories:
    row_cells = table.add_row().cells
    row_cells[0].text = category
    row_cells[1].text = views

doc.save('VIEWS_DOCUMENTATION.docx')
print('Word document created successfully: VIEWS_DOCUMENTATION.docx')
